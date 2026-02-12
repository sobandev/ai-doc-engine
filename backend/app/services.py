import os
import whisper
import docx
from docx.oxml.ns import qn
from pydub import AudioSegment
import requests
import json
import re

# --- CONFIG ---
FFMPEG_DIR = os.path.join(
    os.environ.get("LOCALAPPDATA", ""),
    "Microsoft", "WinGet", "Packages",
    "Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe",
    "ffmpeg-8.0.1-full_build", "bin"
)
if os.path.isdir(FFMPEG_DIR):
    os.environ["PATH"] = FFMPEG_DIR + os.pathsep + os.environ.get("PATH", "")

GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = "llama-3.3-70b-versatile"

# Global model instance to avoid reloading every request
_whisper_model = None

def get_whisper_model():
    global _whisper_model
    if _whisper_model is None:
        print("Loading Whisper model...")
        _whisper_model = whisper.load_model("base")
    return _whisper_model

def transcribe_audio(file_path: str) -> str:
    """Transcribe audio file to text using Whisper."""
    # Convert to wav if needed
    if not file_path.lower().endswith('.wav'):
        new_path = file_path.rsplit('.', 1)[0] + '.wav'
        audio = AudioSegment.from_file(file_path)
        audio.export(new_path, format="wav")
        file_path = new_path
    
    model = get_whisper_model()
    result = model.transcribe(file_path)
    return result["text"].strip()

def extract_placeholders(template_path: str) -> list:
    """Extract placeholders from a docx template."""
    doc = docx.Document(template_path)
    placeholders = set()
    pattern = re.compile(r'\[([^\[\]]+)\]')
    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        placeholders.update(matches)
    return sorted(list(placeholders))

def get_template_text(docx_path: str) -> str:
    """Read full text from docx to give AI context."""
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    # Also read tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = " ".join([p.text for p in cell.paragraphs if p.text.strip()])
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                full_text.append(" | ".join(row_text))
    return "\n".join(full_text)

def parse_transcript_with_ai(transcript: str, placeholders: list, api_key: str, template_text: str = "") -> dict:
    """Use Groq API to parse transcript into structured data."""
    if not placeholders:
        return {}
        
    placeholder_list = "\n".join(f"- {p}" for p in placeholders)
    
    system_prompt = """You are an expert medical scribe AI. Your job is to parse a raw doctor's audio transcription and extract structured information to fill in a medical consultation note template.

RULES:
1. Correct medical errors (spelling, terminology).
2. Be concise but professional.
3. If info is missing, write "Not mentioned".
4. For date, use today's date if not mentioned (YYYY-MM-DD).
5. CRITICAL: Return a FLAT JSON object. All values MUST be strings. Do not use nested objects or arrays.
6. If a value is a list (e.g., findings), join them with commas.
7. If the template contains numbered fields (e.g., Medicine 1, Medicine 2), distribute the items found in the transcript into these fields sequentially.
8. If a field asks for "BP, Pulse, Temp", combine them into one string (e.g., "BP: 120/80, Pulse: 80").
9. Use the TEMPLATE STRUCTURE below to understand the format, tone, and style.
10. If the template contains example/filled values (e.g., "Patient: John Doe"), treat “John Doe” as an EXAMPLE. Extract the NEW value from the transcription (e.g., "Patient: Jane Smith"). Do not copy the example values unless they are static text."""

    user_prompt = f"""Raw Transcription:
{transcript}

TEMPLATE STRUCTURE (Context):
{template_text}

Extract data for these SPECIFIC fields (placeholders):
{placeholder_list}

Return ONLY a valid, flat JSON object mapping field names to string values."""

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": GROQ_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.1,
        "response_format": {"type": "json_object"}
    }
    
    response = requests.post(GROQ_API_URL, headers=headers, json=payload, timeout=60)
    response.raise_for_status()
    return response.json()['choices'][0]['message']['content']

def infer_and_fill_template(transcript: str, template_text: str, api_key: str) -> dict:
    """
    For templates without placeholders:
    1. Identify field:original_value pairs from template_text.
    2. Extract new_value from transcript.
    3. Return { "field_name": { "original": "...", "new": "..." } }
    """
    system_prompt = """You are an expert document analyzer.
The user provides a FILLED template text (e.g. from a previous patient) and a new transcription.
Your job is to:
1. Identify the variable fields in the template (e.g. "Patient Name: Mr. Ali" -> Field: "Patient Name", Original: "Mr. Ali").
2. Extract the NEW value for that field from the transcription (e.g. New: "Jane Doe").
3. Return a JSON object where keys are Field Names, and values are objects containing:
   - "original": The EXACT text value currently in the template (to be replaced).
   - "new": The new value extracted from the transcript.
4. If a "new" value is not mentioned in the transcript, set "new" to "Not mentioned".
5. EXCEPTION: If the field is a Date or Review Date, set "new" to today's date (YYYY-MM-DD) if not mentioned. Do not use "Not mentioned" for dates.
6. IMPORTANT: The "original" value must be unique enough to be replaced safely. If the value is generic (e.g. "No"), include surrounding context if possible or skip.
7. Return purely valid JSON."""

    user_prompt = f"""Template Context (Filled):
{template_text}

New Transcription:
{transcript}

Return JSON mapping fields to original/new values."""

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": GROQ_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.1,
        "response_format": {"type": "json_object"}
    }
    
    response = requests.post(GROQ_API_URL, headers=headers, json=payload, timeout=60)
    response.raise_for_status()
    return response.json()['choices'][0]['message']['content']

def generate_filled_docx(template_path: str, filled_data: dict, output_path: str, replacements: dict = None):
    """Fill template docx with data. Supports both placeholder replacement and direct text replacement."""
    doc = docx.Document(template_path)
    
    clean_data = {k.strip(): str(v) for k, v in filled_data.items()}

    if replacements:
        # Smart Replacement Mode for non-placeholder templates
        for field, meta in replacements.items():
            original = meta.get("original")
            if not original:
                continue
                
            # Get latest value from filled_data, fallback to 'new' from AI
            new_val = clean_data.get(field, meta.get("new", ""))
            
            for para in doc.paragraphs:
                if original in para.text:
                    para.text = para.text.replace(original, new_val)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if original in para.text:
                                para.text = para.text.replace(original, new_val)
    
    # Always try placeholder replacement too (hybrid support)
    for para in doc.paragraphs:
        full_text = para.text
        if '[' not in full_text:
            continue
            
        new_text = full_text
        matches = re.findall(r'\[([^\[\]]+)\]', full_text)
        
        changed = False
        for m in matches:
            m_clean = m.strip()
            if m_clean in clean_data:
                bracket = f"[{m}]"
                if bracket in new_text:
                    new_text = new_text.replace(bracket, clean_data[m_clean])
                    changed = True
        
        if not changed:
            continue
            
        # Nuclear replacement: Clear XML runs and rebuild
        for r in para._element.findall(qn('w:r')):
            para._element.remove(r)
            
        # Try to preserve bold label: "Label: Value"
        colon_pos = new_text.find(': ')
        if colon_pos >= 0:
            label = new_text[:colon_pos + 2]
            value = new_text[colon_pos + 2:]
            run_label = para.add_run(label)
            run_label.bold = True
            run_val = para.add_run(value)
            run_val.bold = False
        else:
            para.add_run(new_text)
            
    doc.save(output_path)
    return output_path
